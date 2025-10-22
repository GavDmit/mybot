import logging
import os
import traceback
from datetime import datetime
from docx import Document
from telegram import Update, ReplyKeyboardMarkup, ReplyKeyboardRemove
from telegram.ext import (
    Application,
    CommandHandler,
    MessageHandler,
    ContextTypes,
    ConversationHandler,
    filters,
)

# Включаем логирование
logging.basicConfig(
    format='%(asctime)s - %(name)s - %(levelname)s - %(message)s',
    level=logging.INFO
)
logger = logging.getLogger(__name__)

# Состояния диалога
(
    ENTER_TITLE,
    ENTER_STAGE_NAME,
    ENTER_START_DATE,
    ENTER_END_DATE,
    ENTER_GOAL,
    ENTER_ACTIVITIES,
    CONFIRM_OR_ADD_MORE
) = range(7)

# Хранилище данных пользователя (в реальном проекте — база данных)
user_data_store = {}

# Команда /start
async def start(update: Update, context: ContextTypes.DEFAULT_TYPE):
    user_id = update.effective_user.id
    user_data_store[user_id] = {
        "title": "",
        "stages": []
    }
    await update.message.reply_text(
        "Привет! Давай создадим календарно-тактический план.\n"
        "Сначала введи название плана:"
    )
    return ENTER_TITLE

# Ввод названия плана
async def enter_title(update: Update, context: ContextTypes.DEFAULT_TYPE):
    user_id = update.effective_user.id
    user_data_store[user_id]["title"] = update.message.text.strip()
    await update.message.reply_text("Название сохранено. Теперь введи название первого этапа:")
    return ENTER_STAGE_NAME

# Ввод названия этапа
async def enter_stage_name(update: Update, context: ContextTypes.DEFAULT_TYPE):
    context.user_data["current_stage"] = {
        "name": update.message.text.strip(),
        "start": "",
        "end": "",
        "goal": "",
        "activities": ""
    }
    await update.message.reply_text("Введите дату начала этапа (в формате ДД.ММ.ГГГГ):")
    return ENTER_START_DATE

# Ввод даты начала
async def enter_start_date(update: Update, context: ContextTypes.DEFAULT_TYPE):
    date_str = update.message.text.strip()
    try:
        datetime.strptime(date_str, "%d.%m.%Y")
        context.user_data["current_stage"]["start"] = date_str
        await update.message.reply_text("Введите дату окончания этапа (в формате ДД.ММ.ГГГГ):")
        return ENTER_END_DATE
    except ValueError:
        await update.message.reply_text("Неверный формат даты. Попробуйте снова (ДД.ММ.ГГГГ):")
        return ENTER_START_DATE

# Ввод даты окончания
async def enter_end_date(update: Update, context: ContextTypes.DEFAULT_TYPE):
    date_str = update.message.text.strip()
    try:
        datetime.strptime(date_str, "%d.%m.%Y")
        context.user_data["current_stage"]["end"] = date_str
        await update.message.reply_text("Введите цель этапа:")
        return ENTER_GOAL
    except ValueError:
        await update.message.reply_text("Неверный формат даты. Попробуйте снова (ДД.ММ.ГГГГ):")
        return ENTER_END_DATE

# Ввод цели
async def enter_goal(update: Update, context: ContextTypes.DEFAULT_TYPE):
    context.user_data["current_stage"]["goal"] = update.message.text.strip()
    await update.message.reply_text("Введите мероприятия (можно несколько, через запятую или перечислите):")
    return ENTER_ACTIVITIES

# Ввод мероприятий
async def enter_activities(update: Update, context: ContextTypes.DEFAULT_TYPE):
    context.user_data["current_stage"]["activities"] = update.message.text.strip()
    user_id = update.effective_user.id
    user_data_store[user_id]["stages"].append(context.user_data["current_stage"].copy())

    reply_keyboard = [["Добавить ещё этап", "Завершить и экспортировать"]]
    await update.message.reply_text(
        "Этап добавлен! Что дальше?",
        reply_markup=ReplyKeyboardMarkup(reply_keyboard, one_time_keyboard=True)
    )
    return CONFIRM_OR_ADD_MORE

# Обработка выбора: добавить ещё или завершить
async def confirm_or_add_more(update: Update, context: ContextTypes.DEFAULT_TYPE):
    text = update.message.text.strip()
    if "ещё" in text.lower():
        await update.message.reply_text("Введите название следующего этапа:", reply_markup=ReplyKeyboardRemove())
        return ENTER_STAGE_NAME
    else:
        await export_plan(update, context)
        return ConversationHandler.END

# Экспорт в Word с обработкой ошибок
async def export_plan(update: Update, context: ContextTypes.DEFAULT_TYPE):
    user_id = update.effective_user.id
    data = user_data_store.get(user_id)
    if not data or not data["stages"]:
        await update.message.reply_text("Нет данных для экспорта.", reply_markup=ReplyKeyboardRemove())
        return

    try:
        doc = Document()
        doc.add_heading(data["title"], 0)

        for i, stage in enumerate(data["stages"], 1):
            doc.add_heading(f"Этап {i}: {stage['name']}", level=1)
            doc.add_paragraph(f"Сроки: с {stage['start']} по {stage['end']}")
            doc.add_paragraph(f"Цель: {stage['goal']}")
            doc.add_paragraph(f"Мероприятия: {stage['activities']}")

        filename = f"Календарно-тактический_план_{user_id}.docx"
        doc.save(filename)

        await update.message.reply_document(
            document=open(filename, 'rb'),
            caption="Ваш календарно-тактический план готов!",
            reply_markup=ReplyKeyboardRemove()
        )

        # Опционально: удаляем файл после отправки
        # os.remove(filename)

    except Exception as e:
        error_msg = f"Ошибка при экспорте: {str(e)}"
        print(error_msg)
        traceback.print_exc()
        await update.message.reply_text(
            "Произошла ошибка при создании файла. Попробуйте снова.",
            reply_markup=ReplyKeyboardRemove()
        )

# Отмена
async def cancel(update: Update, context: ContextTypes.DEFAULT_TYPE):
    await update.message.reply_text("Создание плана отменено.", reply_markup=ReplyKeyboardRemove())
    return ConversationHandler.END

# Основная функция
def main():
    token = os.environ["BOT_TOKEN"]
    application = Application.builder().token(token).build()

    conv_handler = ConversationHandler(
        entry_points=[CommandHandler("start", start)],
        states={
            ENTER_TITLE: [MessageHandler(filters.TEXT & ~filters.COMMAND, enter_title)],
            ENTER_STAGE_NAME: [MessageHandler(filters.TEXT & ~filters.COMMAND, enter_stage_name)],
            ENTER_START_DATE: [MessageHandler(filters.TEXT & ~filters.COMMAND, enter_start_date)],
            ENTER_END_DATE: [MessageHandler(filters.TEXT & ~filters.COMMAND, enter_end_date)],
            ENTER_GOAL: [MessageHandler(filters.TEXT & ~filters.COMMAND, enter_goal)],
            ENTER_ACTIVITIES: [MessageHandler(filters.TEXT & ~filters.COMMAND, enter_activities)],
            CONFIRM_OR_ADD_MORE: [MessageHandler(filters.Regex("^(Добавить ещё этап|Завершить и экспортировать)$"), confirm_or_add_more)],
        },
        fallbacks=[CommandHandler("cancel", cancel)]
    )

    application.add_handler(conv_handler)
    application.run_polling()

if __name__ == '__main__':
    main()